"""
Document extraction — turn an uploaded file (PDF, Word, Excel, CSV, text) into
plain text the rest of the app can use: ask questions about it, or generate a
Study Kit from it.

Scope, honestly: this extracts the TEXT layer. Text PDFs, .docx, .xlsx, .csv,
.txt all work. Scanned/photo PDFs (no text layer) and images/diagrams inside a
file need OCR / a vision model — not handled here; those are a later phase.
Everything used is pure-python and free (no OCR service, no paid API).
"""
from __future__ import annotations

import csv
import io

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File

from app.auth import get_current_user
from app.models import User
from app.security.ratelimit import rate_limit

router = APIRouter(prefix="/api/files", tags=["files"])

MAX_BYTES = 15 * 1024 * 1024   # 15 MB cap
MAX_CHARS = 24000              # keep the extracted text within model context


def _extract_pdf(data: bytes) -> tuple[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    return "\n\n".join(parts), len(reader.pages)


def _extract_docx(data: bytes) -> tuple[str, int]:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    # tables too
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines), 0


def _extract_xlsx(data: bytes) -> tuple[str, int]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            vals = [str(c) for c in row if c is not None]
            if vals:
                out.append(" | ".join(vals))
    return "\n".join(out), len(wb.worksheets)


def _extract_csv(data: bytes) -> tuple[str, int]:
    text = data.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    return "\n".join(" | ".join(r) for r in rows if any(r)), 0


@router.post("/extract")
async def extract_file(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    _: None = Depends(rate_limit(20, 60)),
):
    """Extract plain text from an uploaded document. Returns {text, chars,
    pages, name, truncated}."""
    name = (file.filename or "file").lower()
    data = await file.read()
    if not data:
        raise HTTPException(400, "Empty file.")
    if len(data) > MAX_BYTES:
        raise HTTPException(413, "File too large (max 15 MB).")

    try:
        if name.endswith(".pdf"):
            text, pages = _extract_pdf(data)
        elif name.endswith(".docx"):
            text, pages = _extract_docx(data)
        elif name.endswith((".xlsx", ".xlsm")):
            text, pages = _extract_xlsx(data)
        elif name.endswith(".csv"):
            text, pages = _extract_csv(data)
        elif name.endswith((".txt", ".md")):
            text, pages = data.decode("utf-8", errors="replace"), 0
        elif name.endswith(".doc"):
            raise HTTPException(415, "Old .doc isn't supported — save as .docx and try again.")
        else:
            raise HTTPException(415, "Unsupported file type. Use PDF, DOCX, XLSX, CSV, or TXT.")
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(422, f"Couldn't read that file — it may be scanned/protected. ({type(exc).__name__})")

    text = (text or "").strip()
    if not text:
        raise HTTPException(
            422,
            "No text found. If this is a scanned PDF or images, text extraction "
            "can't read it (that needs OCR, coming later).",
        )
    truncated = len(text) > MAX_CHARS
    return {
        "text": text[:MAX_CHARS],
        "chars": len(text),
        "pages": pages,
        "name": file.filename,
        "truncated": truncated,
    }
